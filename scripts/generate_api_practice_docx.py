#!/usr/bin/env python3
"""Generate the formatted Word API practice manual from Markdown."""

from __future__ import annotations

import argparse
import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEAL = "008C86"
LIGHT_BLUE = "E8EEF5"
LIGHT_TEAL = "E8F5F3"
LIGHT_GRAY = "F4F6F8"
MID_GRAY = "6B7280"
BLACK = "1F2937"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, name="Calibri", east_asia="Microsoft YaHei", size=None,
                 color=None, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_shading(paragraph, fill):
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcpr = tc.get_or_add_tcPr()
    mar = tcpr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tcpr.append(mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    trpr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    trpr.append(header)


def set_table_borders(table, color="B8C3CF", size="4"):
    tblpr = table._tbl.tblPr
    borders = tblpr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblpr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblpr = table._tbl.tblPr

    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(sum(widths_dxa)))
    tblw.set(qn("w:type"), "dxa")

    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tblind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(width))
            tcw.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)


def text_weight(value):
    score = 0
    for char in re.sub(r"[`*_]", "", value):
        score += 2 if ord(char) > 127 else 1
    return max(4, min(score, 60))


def calculate_widths(rows):
    cols = len(rows[0])
    weights = []
    for idx in range(cols):
        longest = max(text_weight(row[idx]) if idx < len(row) else 4 for row in rows)
        header = text_weight(rows[0][idx])
        weights.append(max(longest, header * 1.2, 7))
    minimums = {2: 1100, 3: 900, 4: 760, 5: 650}
    minimum = minimums.get(cols, 600)
    widths = [minimum] * cols
    remaining = TABLE_WIDTH_DXA - minimum * cols
    total = sum(weights)
    for idx, weight in enumerate(weights):
        widths[idx] += int(remaining * weight / total)
    widths[-1] += TABLE_WIDTH_DXA - sum(widths)
    return widths


def parse_inline(paragraph, text, default_size=10.5, default_color=BLACK):
    text = text.replace("  ", " ")
    token_re = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^\)]+\)|\*[^*]+\*)")
    pos = 0
    for match in token_re.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, size=default_size, color=default_color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=default_size, color=default_color, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", east_asia="Microsoft YaHei", size=max(8.5, default_size - 0.5), color=DARK_BLUE)
        elif token.startswith("["):
            label, url = re.match(r"\[([^\]]+)\]\(([^\)]+)\)", token).groups()
            shown = label if url.startswith("#") else f"{label} ({url})"
            run = paragraph.add_run(shown)
            set_run_font(run, size=default_size, color=BLUE)
            run.underline = True
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=default_size, color=default_color, italic=True)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=default_size, color=default_color)


def add_field(paragraph, instruction, placeholder=""):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 11.5, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("微招系统完整接口测试文档")
    set_run_font(run, size=8.5, color=MID_GRAY, bold=True)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("第 ")
    set_run_font(run, size=8.5, color=MID_GRAY)
    add_field(p, " PAGE ", "1")
    run = p.add_run(" 页")
    set_run_font(run, size=8.5, color=MID_GRAY)


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(118)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("WEIB · API PRACTICE MANUAL")
    set_run_font(run, size=10, color=TEAL, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("微招系统完整接口测试文档")
    set_run_font(run, size=28, color=DARK_BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(36)
    run = p.add_run("Apifox / Postman / cURL 学习与练习手册")
    set_run_font(run, size=14, color=MID_GRAY)

    for label, value in (
        ("适用环境", "http://superorange.top"),
        ("文档版本", "2.0"),
        ("覆盖范围", "65 个业务接口 + 24 个页面路由"),
        ("生成日期", "2026-07-10"),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(f"{label}  ")
        set_run_font(r, size=10.5, color=MID_GRAY, bold=True)
        r = p.add_run(value)
        set_run_font(r, size=10.5, color=BLACK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(75)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("从登录开始，逐项写清参数、认证、示例、响应和测试点")
    set_run_font(run, size=10, color=TEAL, italic=True)
    doc.add_page_break()


def add_toc(doc):
    p = doc.add_paragraph("目录", style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(10)
    r = note.add_run("以下页码已按本次生成结果校对；也可使用 Word 导航窗格按标题跳转。")
    set_run_font(r, size=9, color=MID_GRAY, italic=True)

    entries = [
        ("快速找到登录接口", 3),
        ("1. 测试准备", 4),
        ("2. 通用约定", 6),
        ("3. 普通用户登录完整练习", 8),
        ("4. 管理员登录完整练习", 11),
        ("5. 账号与公共接口", 13),
        ("6. 求职者接口", 16),
        ("7. Boss、公司与职位管理接口", 22),
        ("8. 聊天、文件与通知接口", 27),
        ("9. 管理后台接口", 32),
        ("10. 地图与公共数据接口", 41),
        ("11. 错误、限流与幂等", 42),
        ("12. 练习用例与变量记录表", 44),
        ("13. 页面型路由附录", 46),
    ]
    for title, page in entries:
        line = doc.add_paragraph()
        line.paragraph_format.space_after = Pt(3)
        line.paragraph_format.tab_stops.add_tab_stop(
            Inches(6.1), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
        )
        r = line.add_run(title)
        set_run_font(r, size=10.5, color=BLACK, bold=title.startswith(("3.", "4.")))
        r = line.add_run(f"\t{page}")
        set_run_font(r, size=10.5, color=BLUE, bold=True)
    doc.add_page_break()


def add_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    set_table_borders(table)
    widths = calculate_widths(rows)
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for ridx, row in enumerate(rows):
        for cidx, value in enumerate(row):
            cell = table.cell(ridx, cidx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            if ridx == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            short = len(value) <= 12 and not any(ch in value for ch in "，。；：/")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if (ridx == 0 or short) else WD_ALIGN_PARAGRAPH.LEFT
            parse_inline(p, value, default_size=8.3 if len(rows[0]) >= 5 else 8.8)
            if ridx == 0:
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)


def add_code_block(doc, lines, language=""):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    set_paragraph_shading(p, LIGHT_GRAY)
    code = "\n".join(lines).rstrip()
    run = p.add_run(code or " ")
    set_run_font(run, name="Consolas", east_asia="Microsoft YaHei", size=8.3, color="263238")


def add_callout(doc, lines):
    text = " ".join(line.lstrip("> ").strip() for line in lines).strip()
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    set_paragraph_shading(p, LIGHT_TEAL)
    r = p.add_run("提示：")
    set_run_font(r, size=9.5, color=TEAL, bold=True)
    parse_inline(p, text, default_size=9.5)


def markdown_table(lines, start):
    rows = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        parts = [cell.strip() for cell in lines[idx].strip().strip("|").split("|")]
        rows.append(parts)
        idx += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in rows[1]):
        rows.pop(1)
    width = max(len(row) for row in rows)
    rows = [row + [""] * (width - len(row)) for row in rows]
    return rows, idx


def build_doc(markdown_path: Path, output_path: Path):
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_styles(doc)
    for section in doc.sections:
        configure_section(section)
    add_cover(doc)
    add_toc(doc)

    in_code = False
    code_language = ""
    code_lines = []
    skipping_manual_toc = False
    idx = 0
    while idx < len(lines):
        raw = lines[idx]
        stripped = raw.strip()

        if stripped.startswith("# 微招系统完整接口测试文档") or stripped.startswith("> 文档版本") \
                or stripped.startswith("> 适用环境") or stripped.startswith("> 编写日期") \
                or stripped.startswith("> 用途"):
            idx += 1
            continue

        if stripped == "## 目录":
            skipping_manual_toc = True
            idx += 1
            continue
        if skipping_manual_toc:
            if stripped == "---":
                skipping_manual_toc = False
            idx += 1
            continue

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines, code_language)
                in_code = False
                code_lines = []
            else:
                in_code = True
                code_language = stripped[3:].strip()
            idx += 1
            continue
        if in_code:
            code_lines.append(raw)
            idx += 1
            continue

        if stripped.startswith("|"):
            rows, idx = markdown_table(lines, idx)
            add_table(doc, rows)
            continue

        if stripped.startswith(">"):
            block = []
            while idx < len(lines) and lines[idx].strip().startswith(">"):
                block.append(lines[idx])
                idx += 1
            add_callout(doc, block)
            continue

        if stripped.startswith("<a id=") or stripped == "---":
            idx += 1
            continue

        heading = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading:
            hashes, title = heading.groups()
            title = title.replace("`", "")
            level = len(hashes) - 1
            p = doc.add_paragraph(style=f"Heading {min(level, 3)}")
            if level == 1 and re.match(r"\d+\.", title):
                # Word handles page-break-before without inserting an empty page
                # when the heading already starts at the top of a page.
                p.paragraph_format.page_break_before = True
            parse_inline(p, title, default_size={1: 16, 2: 13, 3: 11.5}.get(level, 11.5),
                         default_color={1: BLUE, 2: BLUE, 3: DARK_BLUE}.get(level, DARK_BLUE))
            idx += 1
            continue

        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        number = re.match(r"^\d+\.\s+(.+)$", stripped)
        if bullet or number:
            p = doc.add_paragraph(style="List Bullet" if bullet else "List Number")
            parse_inline(p, (bullet or number).group(1))
            idx += 1
            continue

        if not stripped:
            idx += 1
            continue

        p = doc.add_paragraph()
        p.paragraph_format.keep_together = False
        parse_inline(p, stripped)
        idx += 1

    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")

    doc.core_properties.title = "微招系统完整接口测试文档"
    doc.core_properties.subject = "Weib API practice manual"
    doc.core_properties.author = "Weib Project"
    doc.core_properties.keywords = "API, Apifox, Postman, Spring Boot, Weib"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    args = parser.parse_args()
    build_doc(args.input, args.output)
    print(args.output)


if __name__ == "__main__":
    main()
