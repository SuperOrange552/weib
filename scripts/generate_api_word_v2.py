#!/usr/bin/env python3
"""Generate the concise, complete Weib API testing manual (Word V2)."""

from __future__ import annotations

import argparse
from collections import Counter
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from api_word_v2_catalog import APP_BADGE, DEFAULT_VARIABLES, InterfaceCard, ParameterRow, build_catalog
from api_word_v2_inventory import extract_endpoints


# compact_reference_guide preset, resolved to concrete tokens.
PAGE_WIDTH = Inches(8.5)
PAGE_HEIGHT = Inches(11)
MARGIN = Inches(1)
USABLE_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}
BODY_FONT = "Microsoft YaHei"
LATIN_FONT = "Calibri"
CODE_FONT = "Consolas"
ACCENT = "2E74B5"
DARK = "1F2937"
MUTED = "5B6472"
LIGHT = "E8EEF5"
PALE_BLUE = "F4F8FC"
PALE_GREEN = "E9F7EF"
PALE_ORANGE = "FFF4E5"
BORDER = "C8D3DF"
MODULE_NAMES = {
    "auth": "一、认证与账号",
    "seeker": "二、求职者与通用业务",
    "boss": "三、招聘者业务",
    "community": "四、论坛、投诉与申诉",
    "files": "五、文件与媒体",
    "mobile": "六、Android App 专用接口",
    "admin": "七、管理后台接口",
    "pages": "八、页面路由",
}


def _set_run_font(run, *, name=BODY_FONT, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), LATIN_FONT if name == BODY_FONT else name)
    fonts.set(qn("w:hAnsi"), LATIN_FONT if name == BODY_FONT else name)
    fonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _shade(cell_or_paragraph, fill: str):
    if hasattr(cell_or_paragraph, "_tc"):
        props = cell_or_paragraph._tc.get_or_add_tcPr()
    else:
        props = cell_or_paragraph._p.get_or_add_pPr()
    shd = props.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        props.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in CELL_MARGIN_DXA.items():
        tag = "start" if edge == "start" else "end" if edge == "end" else edge
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths: list[int], indent: int = TABLE_INDENT_DXA):
    if sum(widths) != USABLE_DXA:
        raise ValueError(f"table widths must total {USABLE_DXA}: {widths}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(USABLE_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _set_cell_text(cell, text: str, *, bold=False, color=DARK, size=9.2, align=WD_ALIGN_PARAGRAPH.LEFT, font=BODY_FONT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(str(text))
    _set_run_font(r, name=font, size=size, color=color, bold=bold)


def _set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def _add_page_field(paragraph):
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def _add_toc(paragraph):
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "在 Word 中右键此处并选择“更新域”以刷新目录。"
    run.append(text)
    fld.append(run)
    paragraph._p.append(fld)


def _configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = MARGIN
    section.bottom_margin = MARGIN
    section.left_margin = MARGIN
    section.right_margin = MARGIN
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, ACCENT, 18, 10),
        ("Heading 2", 13, ACCENT, 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def _add_header_footer(section, generated_date: str):
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    r = header.add_run("微招系统接口测试手册 V2")
    _set_run_font(r, size=8.5, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    footer.paragraph_format.space_after = Pt(0)
    r = footer.add_run(f"生成日期：{generated_date}    第 ")
    _set_run_font(r, size=8, color=MUTED)
    _add_page_field(footer)
    r = footer.add_run(" 页")
    _set_run_font(r, size=8, color=MUTED)


def _add_cover(doc: Document, cards: list[InterfaceCard], generated_date: str):
    # editorial_cover: quiet top label, centered title stack, compact facts.
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(34)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("WEIB / 微招")
    _set_run_font(r, size=11, color=MUTED, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("微招系统接口测试手册")
    _set_run_font(r, size=26, color=ACCENT, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    r = p.add_run("精简完整版 V2 · Apifox / Postman 实操版")
    _set_run_font(r, size=14, color=DARK)

    counts = Counter(card.platform for card in cards)
    table = doc.add_table(rows=2, cols=4)
    _set_table_geometry(table, [2340, 2340, 2340, 2340], indent=0)
    values = [
        ("接口总数", str(len(cards))),
        ("业务接口", str(sum(not c.page_route for c in cards))),
        ("APP 专用", str(counts[APP_BADGE])),
        ("页面路由", str(sum(c.page_route for c in cards))),
    ]
    for col, (label, value) in enumerate(values):
        _shade(table.cell(0, col), LIGHT)
        _set_cell_text(table.cell(0, col), label, bold=True, color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(table.cell(1, col), value, bold=True, size=16, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("基准环境  https://superorange.top")
    _set_run_font(r, name=CODE_FONT, size=10.5, color=DARK, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"依据当前 Controller 与 OpenAPI 生成  |  {generated_date}")
    _set_run_font(r, size=9, color=MUTED)
    doc.add_page_break()


def _add_callout(doc: Document, title: str, text: str, fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    _set_table_geometry(table, [USABLE_DXA])
    cell = table.cell(0, 0)
    _shade(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title + "  ")
    _set_run_font(r, size=9.5, color=ACCENT, bold=True)
    r = p.add_run(text)
    _set_run_font(r, size=9.5, color=DARK)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def _add_front_matter(doc: Document, cards: list[InterfaceCard]):
    doc.add_heading("使用说明", level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"标记说明：{APP_BADGE} 表示仅供 Android App 使用；Web/App 通用表示两端可复用；管理后台接口不进入 App。")
    _set_run_font(r, size=9.5, color=ACCENT, bold=True)
    _add_callout(doc, "直接测试", "每个接口均给出完整 URL、权限、请求头、参数示例、可复制请求数据、成功响应、变量来源和关键错误。动态 ID 必须按“变量来源”先获取。")
    _add_callout(doc, "安全提示", "验证码、Cookie、CSRF Token、管理员 Token 和 App Token 均为运行时变量，不在文档中固定保存；不要把生产 Token 写入 Apifox 公共环境。", PALE_ORANGE)

    doc.add_heading("环境变量与测试账号", level=2)
    table = doc.add_table(rows=1, cols=3)
    widths = [1900, 2800, 4660]
    _set_table_geometry(table, widths)
    for i, text in enumerate(("变量", "示例/默认值", "获取与使用方式")):
        _shade(table.cell(0, i), LIGHT)
        _set_cell_text(table.cell(0, i), text, bold=True, color=ACCENT)
    _set_repeat_table_header(table.rows[0])
    variable_rows = [
        ("{{baseUrl}}", DEFAULT_VARIABLES["baseUrl"], "所有接口的基础地址"),
        ("求职者账号", DEFAULT_VARIABLES["seekerUsername"], "统一测试密码：Weib@123456"),
        ("招聘者账号", DEFAULT_VARIABLES["bossUsername"], "统一测试密码：Weib@123456"),
        ("管理员账号", DEFAULT_VARIABLES["adminUsername"], "统一测试密码：Weib@123456；仅管理后台"),
        ("{{captcha}}", "动态 4 位验证码", DEFAULT_VARIABLES["captcha"]),
        ("{{csrfToken}}", "动态", DEFAULT_VARIABLES["csrfToken"]),
        ("{{adminToken}}", "动态", DEFAULT_VARIABLES["adminToken"]),
        ("{{mobileToken}}", "动态", DEFAULT_VARIABLES["mobileToken"]),
        ("{{idempotencyKey}}", "UUID", DEFAULT_VARIABLES["idempotencyKey"]),
    ]
    for variable, value, source in variable_rows:
        cells = table.add_row().cells
        _set_cell_text(cells[0], variable, font=CODE_FONT, size=8.5)
        _set_cell_text(cells[1], value, size=8.8)
        _set_cell_text(cells[2], source, size=8.8)
    _set_table_geometry(table, widths)

    doc.add_heading("认证流程速查", level=2)
    steps = [
        "Web：GET /login → 保存 JSESSIONID → GET /captcha → 读取验证码 → POST /login（同一 Cookie、携带 _csrf）。",
        "App：GET /captcha（保存 Cookie；测试开关开启时可改用 /api/test/captcha 读取明文）→ POST /api/mobile/auth/login → 保存 data.accessToken → 后续使用 Authorization: Bearer {{mobileToken}}。",
        "管理员：POST /api/admin/auth/login → 保存 data.token → 后续使用 Authorization: Bearer {{adminToken}}。",
        "写操作：Session 模式同时携带 X-CSRF-TOKEN；标注 Idempotency-Key 的接口每次新业务生成新 UUID。",
    ]
    for index, text in enumerate(steps, 1):
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        _set_run_font(r, size=10, color=DARK)

    doc.add_heading("接口目录", level=1)
    p = doc.add_paragraph()
    _add_toc(p)
    doc.add_page_break()


def _add_metadata_table(doc: Document, card: InterfaceCard):
    rows = [
        ("功能", card.function),
        ("请求方式与完整地址", f"{card.method} {card.full_url}"),
        ("权限与前置条件", card.permission),
        ("Content-Type", card.content_type),
        ("请求头", "；".join(f"{k}: {v}" for k, v in card.headers) if card.headers else "无额外请求头"),
    ]
    table = doc.add_table(rows=0, cols=2)
    widths = [1700, 7660]
    for label, value in rows:
        cells = table.add_row().cells
        _shade(cells[0], LIGHT)
        _set_cell_text(cells[0], label, bold=True, color=ACCENT, size=8.8)
        _set_cell_text(cells[1], value, font=CODE_FONT if label in ("请求方式与完整地址", "Content-Type") else BODY_FONT, size=8.5 if label == "请求方式与完整地址" else 8.8)
    _set_table_geometry(table, widths)


def _add_parameter_table(doc: Document, rows: tuple[ParameterRow, ...]):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    r = p.add_run("参数表")
    _set_run_font(r, size=9.5, color=ACCENT, bold=True)
    if not rows:
        r = p.add_run("  无")
        _set_run_font(r, size=9, color=MUTED)
        return
    table = doc.add_table(rows=1, cols=6)
    widths = [1540, 1000, 720, 900, 2300, 2900]
    headers = ("参数", "位置", "必填", "类型", "示例", "说明")
    for i, text in enumerate(headers):
        _shade(table.cell(0, i), LIGHT)
        _set_cell_text(table.cell(0, i), text, bold=True, color=ACCENT, size=8.3, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        values = (row.name, row.location, row.required, row.type, row.example, row.description or "—")
        for i, value in enumerate(values):
            _set_cell_text(cells[i], value, font=CODE_FONT if i in (0, 3, 4) else BODY_FONT, size=7.9, align=WD_ALIGN_PARAGRAPH.CENTER if i in (1, 2, 3) else WD_ALIGN_PARAGRAPH.LEFT)
    _set_table_geometry(table, widths)


def _add_code_block(doc: Document, title: str, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title)
    _set_run_font(r, size=9.5, color=ACCENT, bold=True)
    code = doc.add_paragraph()
    _shade(code, "F3F5F7")
    code.paragraph_format.left_indent = Inches(0.12)
    code.paragraph_format.right_indent = Inches(0.12)
    code.paragraph_format.space_before = Pt(0)
    code.paragraph_format.space_after = Pt(4)
    code.paragraph_format.line_spacing = 1.05
    r = code.add_run(text)
    _set_run_font(r, name=CODE_FONT, size=7.7, color=DARK)


def _add_interface_card(doc: Document, card: InterfaceCard):
    # Interface cards are Heading 2 under each module Heading 1, so Word navigation
    # and accessibility hierarchies remain valid without skipped levels.
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"{card.code}｜{card.platform}｜{card.key}")
    _set_run_font(r, size=11.3, color="1F4D78", bold=True)
    _add_metadata_table(doc, card)
    _add_parameter_table(doc, card.parameters)
    _add_code_block(doc, "可复制请求数据", card.request_sample)
    _add_code_block(doc, "成功响应示例", card.success_response)
    _add_callout(doc, "变量来源", card.variable_source, PALE_GREEN)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("关键错误：")
    _set_run_font(r, size=8.5, color=MUTED, bold=True)
    r = p.add_run("；".join(card.key_errors))
    _set_run_font(r, size=8.5, color=MUTED)


def _add_appendix(doc: Document):
    doc.add_page_break()
    doc.add_heading("附录：状态码与测试检查清单", level=1)
    table = doc.add_table(rows=1, cols=2)
    widths = [1400, 7960]
    for i, text in enumerate(("状态码", "测试含义")):
        _shade(table.cell(0, i), LIGHT)
        _set_cell_text(table.cell(0, i), text, bold=True, color=ACCENT)
    for code, meaning in (("200", "成功；同时检查业务 code/data"), ("302", "页面跳转或登录后重定向"), ("400", "参数格式、验证码或业务校验失败"), ("401", "未登录、会话过期或 Token 无效"), ("403", "身份/角色/权限不足或账号受限"), ("404", "资源不存在或路径变量错误"), ("409", "重复提交、幂等冲突或状态冲突"), ("413", "上传文件超过 Nginx/后端大小限制"), ("415", "Content-Type 或文件类型不支持"), ("429", "触发验证码/登录/接口限流"), ("500", "服务端异常，结合 Linux 日志排查")):
        cells = table.add_row().cells
        _set_cell_text(cells[0], code, font=CODE_FONT, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(cells[1], meaning)
    _set_repeat_table_header(table.rows[0])
    _set_table_geometry(table, widths)
    _add_callout(doc, "完成标准", "优先验证认证链，再获取动态 ID；写接口使用专用测试数据并记录清理方式；上传接口使用 1 个合法小文件和 1 个超限/非法类型文件；不在测试报告中保存真实 Token。")


def generate_manual(cards: list[InterfaceCard], output: Path, generated_date: str | None = None) -> Path:
    generated_date = generated_date or date.today().isoformat()
    output.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _configure_document(doc)
    _add_header_footer(doc.sections[0], generated_date)
    _add_cover(doc, cards, generated_date)
    _add_front_matter(doc, cards)
    current_module = None
    for card in cards:
        if card.module != current_module:
            first_module = current_module is None
            current_module = card.module
            module_heading = doc.add_heading(MODULE_NAMES[card.module], level=1)
            # A page-break-before heading avoids the blank page that an explicit
            # page-break paragraph can create when the previous card exactly fills a page.
            module_heading.paragraph_format.page_break_before = not first_module
            if card.module == "mobile":
                _add_callout(doc, APP_BADGE, "本章全部为 Android App 专用接口。除登录外，统一携带 Authorization: Bearer {{mobileToken}}。")
            elif card.module == "admin":
                _add_callout(doc, "管理后台", "本章不属于 Android App；登录后使用管理员 JWT Bearer Token。")
            elif card.module == "pages":
                _add_callout(doc, "页面路由", "本章返回 HTML 页面或重定向，主要用于浏览器手工验证，不按 JSON API 断言。")
        _add_interface_card(doc, card)
    _add_appendix(doc)
    properties = doc.core_properties
    properties.title = "微招系统接口测试手册-精简完整版-V2"
    properties.subject = "Weib Web、Android App 与管理后台接口测试"
    properties.author = "Codex"
    properties.keywords = "Weib, API, Postman, Apifox, Android"
    doc.save(output)
    return output


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--openapi", type=Path, default=Path("docs/verification/api-word-v2-openapi.json"))
    parser.add_argument("--output", type=Path, default=Path("docs/微招系统接口测试手册-精简完整版-V2.docx"))
    args = parser.parse_args()
    root = Path(__file__).resolve().parents[1]
    endpoints = extract_endpoints(root / "src/main/java/com/weib/controller")
    cards = build_catalog(endpoints, root / args.openapi)
    output = generate_manual(cards, root / args.output)
    print(f"generated {output} with {len(cards)} interface cards")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
