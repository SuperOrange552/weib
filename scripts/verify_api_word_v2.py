#!/usr/bin/env python3
"""Strict coverage, labelling and preservation checks for the V2 Word manual."""

from __future__ import annotations

import hashlib
import json
import re
from datetime import datetime, timezone
from pathlib import Path

from docx import Document

from api_word_v2_catalog import APP_BADGE, build_catalog
from api_word_v2_inventory import extract_endpoints


MANUAL = Path("docs/微招系统接口测试手册-精简完整版-V2.docx")
OPENAPI = Path("docs/verification/api-word-v2-openapi.json")
BEFORE = Path("docs/verification/api-word-v2-preservation-before.json")
EVIDENCE = Path("docs/verification/api-word-v2-verification.json")
HEADING_KEY_RE = re.compile(r"^[A-Z]+-\d{3}｜.*?｜(GET|POST|PUT|DELETE|PATCH) (/\S*)$")
REQUIRED_CARD_FIELDS = (
    "功能", "请求方式与完整地址", "权限与前置条件", "Content-Type", "请求头",
    "参数表", "可复制请求数据", "成功响应示例", "变量来源", "关键错误",
)


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def _doc_text(doc_path: Path) -> tuple[str, list[str]]:
    doc = Document(doc_path)
    paragraphs: list[str] = []
    all_text: list[str] = []
    for paragraph in doc.paragraphs:
        paragraphs.append(paragraph.text)
        all_text.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text.extend(p.text for p in cell.paragraphs)
    return "\n".join(all_text), paragraphs


def _preservation_changes(root: Path) -> list[str]:
    records = json.loads((root / BEFORE).read_text(encoding="utf-8-sig"))
    changes: list[str] = []
    for record in records:
        path = Path(record["Path"])
        if not path.exists():
            changes.append(f"missing: {path}")
            continue
        stat = path.stat()
        if stat.st_size != int(record["Length"]):
            changes.append(f"length changed: {path}")
        expected_mtime = datetime.fromisoformat(str(record["LastWriteTime"]).replace("Z", "+00:00"))
        actual_mtime = datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc)
        if abs((actual_mtime - expected_mtime).total_seconds()) > 0.01:
            changes.append(f"mtime changed: {path}")
        if _sha256(path) != record["SHA256"]:
            changes.append(f"sha256 changed: {path}")
    return sorted(set(changes))


def verify_manual(root: Path) -> dict:
    endpoints = extract_endpoints(root / "src/main/java/com/weib/controller")
    cards = build_catalog(endpoints, root / OPENAPI)
    text, paragraphs = _doc_text(root / MANUAL)
    code_keys = {card.key for card in cards}
    documented_keys = {
        f"{match.group(1)} {match.group(2)}"
        for paragraph in paragraphs
        if (match := HEADING_KEY_RE.match(paragraph.strip()))
    }

    app_label_errors: list[str] = []
    for card in cards:
        heading = next((p for p in paragraphs if card.key in p and p.startswith(card.code)), "")
        expected = card.path.startswith("/api/mobile/")
        has_badge = APP_BADGE in heading
        if expected != has_badge:
            app_label_errors.append(card.key)

    incomplete_cards: list[str] = []
    for card in cards:
        values = (
            card.function, card.full_url, card.permission, card.content_type,
            card.request_sample, card.success_response, card.variable_source,
        )
        if not all(value and value.strip() for value in values) or not card.key in text:
            incomplete_cards.append(card.key)
    missing_global_labels = [label for label in REQUIRED_CARD_FIELDS if label not in text]
    if missing_global_labels:
        incomplete_cards.extend(f"GLOBAL:{label}" for label in missing_global_labels)

    changed_old_files = _preservation_changes(root)
    failures = (
        code_keys - documented_keys,
        documented_keys - code_keys,
        app_label_errors,
        incomplete_cards,
        changed_old_files,
    )
    result = {
        "status": "PASS" if not any(failures) else "INCOMPLETE",
        "controllerEndpointCount": len(endpoints),
        "businessEndpointCount": sum(not item.page_route for item in endpoints),
        "pageEndpointCount": sum(item.page_route for item in endpoints),
        "documentedEndpointCount": len(documented_keys),
        "appEndpointCount": sum(item.path.startswith("/api/mobile/") for item in endpoints),
        "missingEndpoints": sorted(code_keys - documented_keys),
        "unknownEndpoints": sorted(documented_keys - code_keys),
        "mislabelledAppEndpoints": sorted(app_label_errors),
        "incompleteCards": sorted(set(incomplete_cards)),
        "changedOldFiles": changed_old_files,
        "requiredCardFields": list(REQUIRED_CARD_FIELDS),
    }
    evidence_path = root / EVIDENCE
    evidence_path.parent.mkdir(parents=True, exist_ok=True)
    existing = {}
    if evidence_path.exists():
        try:
            existing = json.loads(evidence_path.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError):
            existing = {}
    existing.update(result)
    evidence_path.write_text(json.dumps(existing, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    return result


def main() -> int:
    root = Path(__file__).resolve().parents[1]
    result = verify_manual(root)
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0 if result["status"] == "PASS" else 1


if __name__ == "__main__":
    raise SystemExit(main())
